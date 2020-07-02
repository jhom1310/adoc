from docx import *
import re

def sub_palavra(doc_obj, regex , replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
             # Loop adicionado ao trabalho com execuções (strings com o mesmo estilo)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

def sub_tabela(doc_obj, regex, replace):
    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                sub_palavra(cell, regex , replace)
'''
regex = re.compile(r"Aleff") # Palavra que vai ser substituida 
replace1 = r"lukinha" # Palavra para substituição
document = Document('texto.docx') # documento
sub_palavra(document,regex1,replace1)
document.save('testepython.docx')
'''